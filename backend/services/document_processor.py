"""Document text extraction — PyMuPDF for PDFs, python-docx for DOCX."""

from pathlib import Path

import fitz  # PyMuPDF
from docx import Document


def extract_text(file_path: str, mime_type: str) -> str:
    """Extract text from a PDF or DOCX file based on MIME type."""
    if mime_type == "application/pdf":
        return _extract_pdf_text(file_path)
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_docx_text(file_path)
    else:
        raise ValueError(f"Unsupported MIME type: {mime_type}")


def _extract_pdf_text(file_path: str) -> str:
    """Extract text from all pages of a PDF using PyMuPDF."""
    doc = fitz.open(file_path)
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return "\n".join(text_parts)


def _extract_docx_text(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    doc = Document(file_path)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    return "\n".join(paragraphs)
