"""
Workflow simulation script.
Generates a mock CV Word document (.docx), uploads it to the running backend server,
and checks if text extraction, Normalization, and Auto-evaluation operate successfully.
"""

import os
import sys
import time
import requests
from docx import Document

# Base API URL
API_URL = "http://localhost:8000/api"

def create_mock_cv(filepath: str):
    """Generate a high-quality mock CV Word document for testing."""
    doc = Document()
    doc.add_heading("DR. RAHUL KUMAR", level=0)
    
    doc.add_heading("Personal Details", level=1)
    doc.add_paragraph("Email: rahul.kumar@example.com")
    doc.add_paragraph("Phone: +91 9900112233")
    doc.add_paragraph("Address: Hyderabad, India")
    
    doc.add_heading("Education", level=1)
    doc.add_paragraph("Ph.D. in Computer Science — Completed in 2022")
    doc.add_paragraph("University: IIT Bombay")
    
    doc.add_paragraph("Master of Technology (M.Tech.) in Software Engineering — Completed in 2018")
    doc.add_paragraph("University: BITS Pilani")
    
    doc.add_paragraph("Bachelor of Technology (B.Tech.) in Computer Science — Completed in 2016")
    doc.add_paragraph("University: NIT Trichy")
    
    doc.add_heading("Work Experience", level=1)
    doc.add_paragraph("Assistant Professor at Woxsen University (2022 - Present)")
    doc.add_paragraph("Responsibilities: Researching dynamic evaluation systems and lecturing.")
    
    doc.save(filepath)
    print(f"Mock CV created at {filepath}")


def run_simulation():
    print("=" * 60)
    print("STARTING WORKFLOW SIMULATION")
    print("=" * 60)

    # 1. Clean old test file if exists
    cv_path = "test_cv.docx"
    if os.path.exists(cv_path):
        os.remove(cv_path)

    # 2. Create mock CV document
    create_mock_cv(cv_path)

    # 3. Verify backend health
    try:
        health = requests.get(f"{API_URL}/health")
        print(f"Backend Status: {health.json()['status']}")
    except Exception as e:
        print(f"Error: Backend server is not running on port 8000: {e}")
        return

    # 4. Seed Specialization & Rule (Ensure matching rule exists)
    print("\n[Step 1] Seeding Specialization Mappings...")
    spec_data = {
        "name": "Computer Science",
        "department": "CSE",
        "is_allied": False,
        "aliases": ["CS", "Computer Science", "CSE"]
    }
    requests.post(f"{API_URL}/specializations", json=spec_data)

    print("\n[Step 2] Seeding Criteria 1 Rule...")
    rule_data = {
        "name": "Criteria 1 (CSE Core)",
        "description": "Must have B.Tech, M.Tech, and completed PhD in CSE",
        "department": "CSE",
        "position": "Assistant Professor",
        "priority": 1,
        "is_active": True,
        "logic_operator": "AND",
        "conditions": [
            {"field": "ug_degree", "operator": "equals", "value": "B.Tech.", "order": 0},
            {"field": "pg_degree", "operator": "equals", "value": "M.Tech.", "order": 1},
            {"field": "phd_status", "operator": "equals", "value": "completed", "order": 2}
        ]
    }
    rule_res = requests.post(f"{API_URL}/rules", json=rule_data)
    print(f"Rule created: {rule_res.json()['name']} (Priority: {rule_res.json()['priority']})")

    # 5. Upload Resume
    print("\n[Step 3] Ingesting Resume...")
    with open(cv_path, "rb") as f:
        files = {"files": (cv_path, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        data = {"batch_name": "Simulation Run - Batch 1"}
        upload_res = requests.post(f"{API_URL}/resumes/upload", files=files, data=data)
        
    resumes = upload_res.json()["resumes"]
    if not resumes:
        print("Upload failed: No resumes returned")
        return
        
    resume = resumes[0]
    candidate_id = resume["candidate_id"]
    print(f"File processed: {resume['original_filename']}")
    print(f"Status: {resume['status']}")
    print(f"Candidate ID generated: {candidate_id}")

    # 6. Check Candidate Eligibility Status (Auto-Evaluation validation)
    print("\n[Step 4] Validating Auto-Evaluation...")
    candidate_res = requests.get(f"{API_URL}/candidates/{candidate_id}")
    candidate = candidate_res.json()
    print(f"Candidate Name: {candidate['name']}")
    print(f"Extracted Email: {candidate['email']}")
    print(f"PhD Status: {candidate['phd_status']}")
    print(f"Eligibility Status: {candidate['eligibility_status']}")
    print(f"Review Reason: {candidate['review_reason'] or 'None'}")

    assert candidate["eligibility_status"] == "eligible", "Simulation Failed: Candidate should be eligible"
    print("\nSUCCESS: End-to-end workflow successfully completed and verified! Candidate is marked as ELIGIBLE automatically upon resume upload.")

    # Clean up
    if os.path.exists(cv_path):
        os.remove(cv_path)


if __name__ == "__main__":
    run_simulation()
