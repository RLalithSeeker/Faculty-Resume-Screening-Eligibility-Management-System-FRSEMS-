import os
from docx import Document

def create_resume_docx(filepath: str, name: str, email: str, phone: str, education: list, experience: list, publications: list = None, awards: list = None):
    doc = Document()
    
    # Header
    doc.add_heading(name, level=0)
    
    # Contact Info
    doc.add_heading("Contact Details", level=1)
    doc.add_paragraph(f"Email: {email}")
    doc.add_paragraph(f"Phone: {phone}")
    doc.add_paragraph("Address: Gachibowli, Hyderabad, Telangana, India")
    
    # Education
    doc.add_heading("Education", level=1)
    for edu in education:
        doc.add_paragraph(edu)
        
    # Experience
    doc.add_heading("Professional Experience", level=1)
    for exp in experience:
        doc.add_paragraph(exp)
        
    # Publications (for realistic detail)
    if publications:
        doc.add_heading("Key Publications & Patents", level=1)
        for pub in publications:
            doc.add_paragraph(pub)
            
    # Awards & Honors
    if awards:
        doc.add_heading("Awards & Honors", level=1)
        for award in awards:
            doc.add_paragraph(award)
        
    doc.save(filepath)
    print(f"Mock resume created: {filepath}")

def generate_all():
    target_dir = os.path.join(os.path.dirname(__file__), "mock_resumes")
    os.makedirs(target_dir, exist_ok=True)
    
    # 1. Eligible CSE candidate
    create_resume_docx(
        os.path.join(target_dir, "dr_rahul_kumar_eligible.docx"),
        "Dr. Rahul Kumar",
        "rahul.kumar@example.com",
        "+91 9900112233",
        [
            "Ph.D. in Computer Science — Completed in 2022 from Indian Institute of Technology (IIT) Bombay",
            "Master of Technology (M.Tech.) in Software Engineering — Completed in 2018 from BITS Pilani",
            "Bachelor of Technology (B.Tech.) in Computer Science — Completed in 2016 from NIT Trichy"
        ],
        [
            "Assistant Professor, Department of CSE, Woxsen University (2022 - Present)",
            "Research Fellow, Department of Computer Science, IIT Bombay (2018 - 2022)",
            "Lecturer, Computer Science Department, NIT Trichy (2016 - 2018)"
        ],
        [
            "Kumar, R. et al. 'Dynamic Rule Evaluation in High-Throughput Screening Engines.' IEEE Transactions on Software Engineering, 2023.",
            "Kumar, R. 'State-transition mapping architectures for relational databases.' ACM SIGMOD, 2021."
        ],
        [
            "Recipient of NIT Trichy Best Young Academician Award, 2018",
            "Winner of IIT Bombay Doctorate Research Fellowship, 2019"
        ]
    )
    
    # 2. Non-eligible candidate (PhD Pursuing instead of Completed)
    create_resume_docx(
        os.path.join(target_dir, "prof_priya_sharma_non_eligible.docx"),
        "Prof. Priya Sharma",
        "priya.sharma@example.com",
        "+91 9988776655",
        [
            "Ph.D. in Computer Science — Pursuing (Thesis submitted) at IIIT Hyderabad",
            "Master of Technology (M.Tech.) in Computer Science — Completed in 2019 from Osmania University",
            "Bachelor of Technology (B.Tech.) in Information Technology — Completed in 2017 from JNTU Hyderabad"
        ],
        [
            "Assistant Professor, Department of IT, Anurag Group of Institutions (2019 - Present)",
            "Teaching Assistant, Computer Science Division, IIIT Hyderabad (2020 - 2023)"
        ],
        [
            "Sharma, P. 'Anomaly Detection in Distributed IoT Graph Systems.' Springer Lecture Notes, 2022."
        ]
    )

    # 3. Manual Review candidate (Missing PG degree in qualifications)
    create_resume_docx(
        os.path.join(target_dir, "dr_amit_patel_manual_review.docx"),
        "Dr. Amit Patel",
        "amit.patel@example.com",
        "+91 9123456789",
        [
            "Ph.D. in Computer Engineering — Completed in 2020 from NIT Warangal",
            "Bachelor of Engineering (B.E.) in Electronics — Completed in 2015 from Mumbai University"
        ],
        [
            "Senior Researcher, AI & Analytics Group, TCS Innovation Labs (2020 - Present)",
            "Research Scholar, Computer Science Department, NIT Warangal (2015 - 2020)"
        ]
    )

    # 4. Allied Specialization candidate (M.Tech in Information Technology)
    create_resume_docx(
        os.path.join(target_dir, "dr_sarah_jones_allied.docx"),
        "Dr. Sarah Jones",
        "sarah.jones@example.com",
        "+91 9876543210",
        [
            "Ph.D. in Information Technology — Completed in 2021 from JNTU Hyderabad",
            "Master of Technology (M.Tech.) in Information Technology — Completed in 2017 from JNTU",
            "Bachelor of Technology (B.Tech.) in Computer Science — Completed in 2015 from NIT Calicut"
        ],
        [
            "Assistant Professor, School of Technology, Mahindra University (2021 - Present)"
        ]
    )

    # 5. Low Experience candidate (0 years experience)
    create_resume_docx(
        os.path.join(target_dir, "dr_john_doe_low_experience.docx"),
        "Dr. John Doe",
        "john.doe@example.com",
        "+91 9000100020",
        [
            "Ph.D. in Computer Science — Completed in 2025 from IIT Delhi",
            "Master of Technology (M.Tech.) in Software Engineering — Completed in 2021 from Delhi Technological University (DTU)",
            "Bachelor of Technology (B.Tech.) in Computer Science — Completed in 2019 from Netaji Subhas University of Technology (NSUT)"
        ],
        [
            "Fresh Doctorate Graduate - No formal teaching or industry experience recorded."
        ]
    )

    # 6. Eligible PhD from Top International School (MIT)
    create_resume_docx(
        os.path.join(target_dir, "dr_rajesh_hegde_eligible_mit.docx"),
        "Dr. Rajesh Hegde",
        "rajesh.hegde@example.com",
        "+91 8011223344",
        [
            "Ph.D. in Computer Science — Completed in 2018 from Massachusetts Institute of Technology (MIT), USA",
            "Master of Technology (M.Tech.) in Computer Science — Completed in 2014 from IIT Delhi",
            "Bachelor of Technology (B.Tech.) in Computer Science — Completed in 2012 from NIT Surathkal"
        ],
        [
            "Associate Professor, School of Computer Science, Woxsen University (2022 - Present)",
            "Assistant Professor, CSE Department, IIT Delhi (2018 - 2022)"
        ],
        [
            "Hegde, R. 'Distributed Consensus Algorithms for Edge Databases.' ACM Transactions on Systems, 2020."
        ]
    )

    # 7. Non-eligible candidate (Ph.D and Degrees in Electrical Engineering - Not CS/IT)
    create_resume_docx(
        os.path.join(target_dir, "dr_ananya_rao_non_eligible_ee.docx"),
        "Dr. Ananya Rao",
        "ananya.rao@example.com",
        "+91 8899001122",
        [
            "Ph.D. in Electrical Engineering — Completed in 2021 from IIT Madras",
            "Master of Technology (M.Tech.) in Power Systems — Completed in 2016 from NIT Trichy",
            "Bachelor of Technology (B.Tech.) in Electrical & Electronics Engineering — Completed in 2014 from JNTU Hyderabad"
        ],
        [
            "Assistant Professor, Department of EE, Osmania University Engineering College (2021 - Present)"
        ]
    )

    # 8. Allied Specialization candidate (Electronics & Communication Engineering)
    create_resume_docx(
        os.path.join(target_dir, "dr_vikram_singh_allied_ece.docx"),
        "Dr. Vikram Singh",
        "vikram.singh@example.com",
        "+91 7011223344",
        [
            "Ph.D. in Electronics and Communication Engineering — Completed in 2019 from NIT Rourkela",
            "Master of Technology (M.Tech.) in VLSI Design — Completed in 2015 from IIIT Bangalore",
            "Bachelor of Technology (B.Tech.) in ECE — Completed in 2013 from JNTU Kakinada"
        ],
        [
            "Assistant Professor, CSE Department, Woxsen University (2021 - Present)",
            "VLSI Design Architect, Intel Technologies, Bangalore (2015 - 2021)"
        ]
    )

    # 9. Manual Review candidate (PhD from Stanford but missing M.Tech record)
    create_resume_docx(
        os.path.join(target_dir, "dr_srinivas_murthy_manual_review_gaps.docx"),
        "Dr. Srinivas Murthy",
        "srinivas.murthy@example.com",
        "+91 7788990011",
        [
            "Ph.D. in Computer Science — Completed in 2022 from Stanford University, California",
            "Bachelor of Technology (B.Tech.) in Computer Science — Completed in 2017 from IIT Madras"
        ],
        [
            "Postdoctoral Researcher, Stanford AI Lab (2022 - Present)",
            "Research Intern, Microsoft Research India (2016 - 2017)"
        ]
    )

    # 10. Highly Experienced Eligible Candidate (12+ Years Experience)
    create_resume_docx(
        os.path.join(target_dir, "dr_meera_nair_eligible_teaching.docx"),
        "Dr. Meera Nair",
        "meera.nair@example.com",
        "+91 9911223344",
        [
            "Ph.D. in Computer Science — Completed in 2012 from IIT Madras",
            "Master of Technology (M.Tech.) in Computer Science — Completed in 2008 from NIT Calicut",
            "Bachelor of Technology (B.Tech.) in CSE — Completed in 2006 from College of Engineering, Trivandrum"
        ],
        [
            "Professor & Head of CSE Department, Woxsen University (2020 - Present)",
            "Associate Professor, CSE Division, IIT Madras (2016 - 2020)",
            "Assistant Professor, Computer Science, NIT Calicut (2012 - 2016)"
        ]
    )

if __name__ == "__main__":
    generate_all()
